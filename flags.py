from docx import Document
from docx.shared import Pt
import datetime
import smtplib
import ssl
from email.message import EmailMessage
import imghdr
import fireDetection

#flags
fire_email = False
started = False
name = "_"
email = "_"

def gen_report():
    time = datetime.datetime.now().strftime("%H:%M  %a,  %b%y")
    # Load the existing .docx file
    doc = Document("reports/report.docx")
    # Define the default font size (e.g., 16 points)
    default_font_size = Pt(14)
    # Set the default font size for the entire document
    doc.styles['Normal'].font.size = default_font_size

    doc.add_paragraph(f"Login Details :                                                 {time}")
    doc.add_paragraph(f"Name: {name}                                        Email: {email}")
    doc.add_paragraph("Detected Events : ->")
    doc.add_paragraph("")

    # Add text based on flags
    if fire_email:
        doc.add_paragraph("Fire was detected!!")
        
        # Sending the email
        time = datetime.datetime.now().strftime("%H:%M  %a,  %b%y")
        email_sender = "dailydiscovery678@gmail.com"
        # email_password = os.environ.get('PY_PASS')
        email_password = "ygdxmbrzfakreasx"
        email_receiver = email
        
        # Set the subject and body of the email
        subject = 'Alert: Fire Detected'
        body = f"""
        Name: {name}  
        Time : {time}
                    
        Detected Events : ->
        Fire was detected!!
        """

        em = EmailMessage()
        em['From'] = email_sender
        em['To'] = email_receiver
        em['Subject'] = subject
        em.set_content(body)

        with open(fireDetection.output_path, 'rb') as f:
            file_data = f.read()
            file_type = imghdr.what(f.name)
            file_name = f.name
                    
        em.add_attachment(file_data, maintype='image', subtype=file_type, filename=file_name)
                    
        # Add SSL (layer of security)
        context = ssl.create_default_context()

        # Log in and send the email
        with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=context) as smtp:
            smtp.login(email_sender, email_password)
            smtp.send_message(em)

    # Save the modified .docx file
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    doc.save(f"reports/report_{timestamp}.docx")

